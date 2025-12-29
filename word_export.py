from __future__ import annotations

from dataclasses import asdict
from datetime import datetime

from scour_calc import D21Result, D22Result, D21_VELOCITY_EXPONENT


def _require_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        return Document, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, Cm, Pt
    except Exception as e:
        raise ImportError("缺少依赖：python-docx（请先 pip install python-docx）") from e


def _ensure_docx_suffix(path: str) -> str:
    p = str(path)
    return p if p.lower().endswith(".docx") else (p + ".docx")


def _fmt(x, nd: int = 6) -> str:
    try:
        v = float(x)
        if abs(v) >= 1e4 or (abs(v) > 0 and abs(v) < 1e-3):
            return f"{v:.{nd}e}"
        return f"{v:.{nd}f}".rstrip("0").rstrip(".")
    except Exception:
        return str(x)


def _build_doc_base():
    Document, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, qn, Cm, Pt = _require_docx()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    body_font_pt = Pt(14)  # 四号
    first_line_indent = Pt(28)  # 约等于2字符

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = body_font_pt

    def apply_body_para_format(p, *, level: int = 0):
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.line_spacing = 1.5
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        if level <= 0:
            pf.first_line_indent = first_line_indent
        else:
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(28 * level)

    def add_title(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.line_spacing = 1.5

    def add_h(text: str):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.line_spacing = 1.5

    def add_line(text: str, *, level: int = 0):
        p = doc.add_paragraph()
        apply_body_para_format(p, level=level)
        p.add_run(text)
        return p

    return doc, add_title, add_h, add_line, Cm


def export_d21_docx(
    *,
    path: str,
    name: str | None,
    inputs: dict,
    result: D21Result,
) -> str:
    path = _ensure_docx_suffix(path)
    doc, add_title, add_h, add_line, Cm = _build_doc_base()

    title_name = (name or "").strip()
    suffix = f" - {title_name}" if title_name else ""
    add_title(f"冲刷深度计算书 - D.2.1 丁坝一般冲刷{suffix}")
    add_line(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_line("计算依据：规范 D.2.1（非淹没丁坝一般冲刷深度）。")

    add_h("1  已知条件")
    add_line(
        "H0={H0} m，d50={d50} m，U={U} m/s，L0={L0} m，B={B} m".format(
            H0=_fmt(inputs.get("H0"), 6),
            d50=_fmt(inputs.get("d50"), 6),
            U=_fmt(inputs.get("U"), 6),
            L0=_fmt(inputs.get("L0"), 6),
            B=_fmt(inputs.get("B"), 6),
        )
    )
    add_line(
        "θ={theta}°，m={m}，k1类型={k1_type}".format(
            theta=_fmt(inputs.get("theta_deg"), 6),
            m=_fmt(inputs.get("m"), 6),
            k1_type=str(inputs.get("k1_type")),
        )
    )

    uc_method = str(inputs.get("uc_method"))
    if uc_method == "手动输入":
        add_line(f"Uc 取值：手动输入，Uc={_fmt(inputs.get('uc_manual'), 6)} m/s")
    else:
        add_line(
            "Uc 取值：{mth}，γs={gs} kN/m³，γ={gw} kN/m³".format(
                mth=uc_method,
                gs=_fmt(inputs.get("gamma_s"), 6),
                gw=_fmt(inputs.get("gamma_w"), 6),
            )
        )

    add_h("2  计算过程")
    add_line(f"速度项指数 a 固定为 {D21_VELOCITY_EXPONENT:.2f}。")
    add_line(f"k1={_fmt(result.k1, 6)}，k2={_fmt(result.k2, 6)}，k3={_fmt(result.k3, 6)}")
    add_line(f"Um={_fmt(result.Um, 6)} m/s，Uc={_fmt(result.Uc, 6)} m/s")

    # 速度项 v = (Um-Uc)/sqrt(g*d50)
    try:
        v_term = (float(result.Um) - float(result.Uc)) / ((9.81 * float(inputs.get("d50"))) ** 0.5)
    except Exception:
        v_term = None
    if v_term is not None:
        add_line(f"v=(Um−Uc)/sqrt(g·d50)={_fmt(v_term, 6)}")

    add_line(f"hs/H0={_fmt(result.hs_over_H0, 6)}")

    add_h("3  计算结果")
    add_line(f"hs={_fmt(result.hs, 6)} m")

    add_h("附  中间量")
    for k, v in asdict(result).items():
        add_line(f"{k} = {_fmt(v, 12)}", level=1)

    # 添加图片附件
    import os
    base_dir = os.path.dirname(__file__)
    img1_path = os.path.join(base_dir, "1.png")
    img2_path = os.path.join(base_dir, "2.png")
    
    if os.path.exists(img1_path):
        add_h("附图1")
        doc.add_picture(img1_path, width=Cm(14))
    
    if os.path.exists(img2_path):
        add_h("附图2")
        doc.add_picture(img2_path, width=Cm(14))

    doc.save(path)
    return path


def export_d22_docx(
    *,
    path: str,
    name: str | None,
    inputs: dict,
    result: D22Result,
) -> str:
    path = _ensure_docx_suffix(path)
    doc, add_title, add_h, add_line, Cm = _build_doc_base()

    title_name = (name or "").strip()
    suffix = f" - {title_name}" if title_name else ""
    add_title(f"冲刷深度计算书 - D.2.2 护岸局部冲刷{suffix}")
    add_line(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_line("计算依据：规范 D.2.2（顺坡及平顺护岸局部冲刷深度）。")

    add_h("1  已知条件")
    add_line(
        "H0={H0} m，U={U} m/s，Uc={Uc} m/s，α={alpha}°，n={n}".format(
            H0=_fmt(inputs.get("H0"), 6),
            U=_fmt(inputs.get("U"), 6),
            Uc=_fmt(inputs.get("Uc"), 6),
            alpha=_fmt(inputs.get("alpha_deg"), 6),
            n=_fmt(inputs.get("n"), 6),
        )
    )

    add_h("2  计算过程")
    add_line(f"η（表 D.2.2）={_fmt(result.eta, 6)}")
    add_line(f"Uep=U·(2η/(1+η))={_fmt(result.Uep, 6)} m/s")
    add_line(f"hs=H0·((Uep/Uc)^n−1)={_fmt(result.hs_local, 6)} m")

    add_h("3  计算结果")
    add_line(f"hs(局部)={_fmt(result.hs_local, 6)} m")

    add_h("附  中间量")
    for k, v in asdict(result).items():
        add_line(f"{k} = {_fmt(v, 12)}", level=1)
    
    # 添加图片附件
    import os
    base_dir = os.path.dirname(__file__)
    img1_path = os.path.join(base_dir, "1.png")
    img2_path = os.path.join(base_dir, "2.png")
    
    if os.path.exists(img1_path):
        add_h("附图1")
        doc.add_picture(img1_path, width=Cm(14))
    
    if os.path.exists(img2_path):
        add_h("附图2")
        doc.add_picture(img2_path, width=Cm(14))

    doc.save(path)
    return path
